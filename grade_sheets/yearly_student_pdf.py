import os
import logging
import pythoncom
from docx import Document
from docx2pdf import convert
from django.conf import settings
from grade_sheets.helpers import get_grade_sheet_data
from pass_and_failed.models import PassFailedStatus
from .models import StudentGradeSheetPDF
from academic_years.models import AcademicYear
from datetime import datetime
from pathlib import Path
from grade_sheets.pdf_utils import replace_placeholders
import time

logger = logging.getLogger(__name__)

def generate_yearly_student_pdf(student_id, level_id, academic_year_id, pass_template=True, conditional=False):
    """Generate PDF for a single student with yearly grades."""
    try:
        output_dir = Path(settings.MEDIA_ROOT) / 'output_gradesheets'
        temp_dir = Path(settings.MEDIA_ROOT) / 'temp'
        output_dir.mkdir(parents=True, exist_ok=True)
        temp_dir.mkdir(parents=True, exist_ok=True)

        # Resolve academic_year name for template selection
        try:
            academic_year_obj = AcademicYear.objects.get(id=academic_year_id)
            academic_year = academic_year_obj.name
        except AcademicYear.DoesNotExist:
            logger.error(f"Invalid academic_year_id: {academic_year_id}")
            return []

        # Override pass_template and conditional based on PassFailedStatus
        try:
            status_obj = PassFailedStatus.objects.get(
                student_id=student_id,
                level_id=level_id,
                academic_year_id=academic_year_id
            )
            pass_template = status_obj.status in ['PASS', 'CONDITIONAL']
            conditional = status_obj.status == 'CONDITIONAL'
            logger.debug(f"Using PassFailedStatus: status={status_obj.status}, pass_template={pass_template}, conditional={conditional}")
        except PassFailedStatus.DoesNotExist:
            logger.warning(f"No PassFailedStatus found for student_id={student_id}, level_id={level_id}, academic_year_id={academic_year_id}")
            pass_template = True  # Fallback
            conditional = False

        student_data = get_grade_sheet_data(student_id, level_id, academic_year_id, is_yearly=True)
        if not student_data or 'name' not in student_data:
            logger.warning(f"No data found for student_id={student_id}, level_id={level_id}, academic_year_id={academic_year_id}")
            return []

        # Determine template based on pass/failed status
        template_name = (
            'yearly_card_conditional.docx' if conditional else
            'yearly_card_pass.docx' if pass_template else
            'yearly_card_failed.docx'
        )
        template_path = Path(settings.MEDIA_ROOT) / 'templates' / template_name
        template_path = template_path.resolve()  # Normalize path for Windows

        # Verify template file
        if not template_path.exists():
            logger.error(f"Template file not found: {template_path}. Ensure it exists in {Path(settings.MEDIA_ROOT) / 'templates'}")
            return []
        if not template_path.is_file():
            logger.error(f"Template path {template_path} is not a file")
            return []
        try:
            with template_path.open('rb') as f:
                f.read(1)  # Test read permission
            logger.debug(f"Template file {template_path} is readable, size={template_path.stat().st_size} bytes, modified={datetime.fromtimestamp(template_path.stat().st_mtime)}")
        except PermissionError:
            logger.error(f"Permission denied for template file: {template_path}")
            return []
        except Exception as e:
            logger.error(f"Cannot access template file {template_path}: {str(e)}")
            return []

        # Load template with python-docx
        try:
            doc = Document(str(template_path))
            logger.debug(f"Loaded template {template_path} with {len(doc.paragraphs)} paragraphs")
        except Exception as e:
            logger.error(f"Error loading template {template_path} with python-docx: {str(e)}")
            # Fallback to periodic template
            fallback_template = Path(settings.MEDIA_ROOT) / 'templates' / 'report_card.docx'
            if fallback_template.exists():
                logger.info(f"Attempting fallback to periodic template: {fallback_template}")
                try:
                    doc = Document(str(fallback_template))
                    logger.debug(f"Loaded fallback template {fallback_template} with {len(doc.paragraphs)} paragraphs")
                except Exception as e:
                    logger.error(f"Error loading fallback template {fallback_template}: {str(e)}")
                    return []
            else:
                logger.error(f"Fallback template {fallback_template} not found")
                return []

        doc = replace_placeholders(doc, student_data)
        student_name = student_data['name'].replace(' ', '_').replace(':', '_').replace('/', '_').replace('\\', '_')
        timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")
        temp_filename = f"temp_yearly_card_{student_name}_{academic_year_id}_{timestamp}.docx"
        docx_path = temp_dir / temp_filename
        pdf_filename = temp_filename.replace('.docx', '.pdf')
        pdf_path = output_dir / pdf_filename

        doc.save(str(docx_path))
        logger.info(f"Saved .docx: {docx_path}")

        # Retry conversion up to 3 times to handle RPC errors
        max_attempts = 3
        for attempt in range(1, max_attempts + 1):
            pythoncom.CoInitialize()
            try:
                convert(str(docx_path), str(pdf_path))
                if pdf_path.exists():
                    logger.info(f"Converted to PDF: {pdf_path}")
                    break
                else:
                    logger.warning(f"PDF not created at {pdf_path} on attempt {attempt}")
                    if attempt < max_attempts:
                        time.sleep(2)  # Wait before retry
            except Exception as e:
                logger.error(f"Conversion attempt {attempt} failed: {str(e)}")
                if attempt < max_attempts:
                    time.sleep(2)  # Wait before retry
            finally:
                pythoncom.CoUninitialize()
        else:
            logger.error(f"Failed to convert to PDF after {max_attempts} attempts")
            return []

        # Check for existing StudentGradeSheetPDF record
        existing_pdf = StudentGradeSheetPDF.objects.filter(
            student_id=student_id,
            level_id=level_id,
            academic_year_id=academic_year_id
        ).first()
        if existing_pdf:
            logger.info(f"Existing PDF record found for student_id={student_id}, level_id={level_id}, academic_year_id={academic_year_id}. Updating.")
            existing_pdf.pdf_path = str(pdf_path)
            existing_pdf.filename = pdf_path.name
            existing_pdf.created_at = datetime.now()
            existing_pdf.save()
        else:
            StudentGradeSheetPDF.objects.create(
                student_id=student_id,
                level_id=level_id,
                academic_year=academic_year_obj,
                pdf_path=str(pdf_path),
                filename=pdf_path.name,
                created_at=datetime.now()
            )
            logger.info(f"Created new StudentGradeSheetPDF for student_id={student_id}, level_id={level_id}, academic_year_id={academic_year_id}")

        try:
            docx_path.unlink()
            logger.info(f"Cleaned up .docx: {docx_path}")
        except Exception as e:
            logger.warning(f"Failed to clean up .docx: {str(e)}")

        return [str(pdf_path)]

    except Exception as e:
        logger.error(f"Error generating yearly student PDF for student_id={student_id}, level_id={level_id}, academic_year_id={academic_year_id}: {str(e)}")
        return []