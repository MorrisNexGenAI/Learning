from docx import Document
import logging

logger = logging.getLogger(__name__)

def replace_placeholders(doc, data):
    """Replace {{name}} in paragraphs and {{s[i].sn}}, {{s[i][key]}} in tables."""
    try:
        # Log document content before replacement
        logger.info("Document paragraphs before replacement:")
        for para in doc.paragraphs:
            logger.info(f"Paragraph: {para.text}")
        logger.info("Tables before replacement:")
        for table in doc.tables:
            for row in table.rows:
                row_text = [cell.text.strip() for cell in row.cells]
                logger.info(f"Row: {row_text}")

        # Map template placeholders to data keys
        period_map = {
            '"1"': '1st',
            '"2"': '2nd',
            '"3"': '3rd',
            '"1s"': '1exam',
            '"1a"': '1a',
            '"4"': '4th',
            '"5"': '5th',
            '"6"': '6th',
            '"2s"': '2exam',
            '"2a"': '2a',
            '"f"': 'f'
        }

        # Replace {{name}} in paragraphs
        for paragraph in doc.paragraphs:
            for run in paragraph.runs:
                if '{{name}}' in run.text:
                    run.text = run.text.replace('{{name}}', str(data.get('name', '')))
                    logger.info(f"Replaced {{name}} with {data.get('name', '')} in paragraph")

        # Replace placeholders in tables
        subjects = data.get('s', [])
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    cell_text = cell.text.strip()
                    # Log cell text for debugging
                    if '{{s' in cell_text:
                        logger.debug(f"Processing cell: {cell_text}")
                    for i in range(9):  # Up to 9 subjects
                        # Replace subject name
                        sn_placeholder = f'{{{{s[{i}].sn}}}}'
                        if sn_placeholder in cell_text:
                            value = subjects[i].get('sn', '-') if i < len(subjects) else '-'
                            cell.text = cell.text.replace(sn_placeholder, str(value))
                            logger.info(f"Replaced {sn_placeholder} with {value} in table cell")
                        # Replace grades
                        for template_key, db_key in period_map.items():
                            grade_placeholder = f'{{{{s[{i}][{template_key}]}}}}'
                            if grade_placeholder in cell_text:
                                value = subjects[i].get(db_key, '-') if i < len(subjects) else '-'
                                cell.text = cell.text.replace(grade_placeholder, str(value))
                                logger.info(f"Replaced {grade_placeholder} with {value} in table cell")

        # Log document content after replacement
        logger.info("Document paragraphs after replacement:")
        for para in doc.paragraphs:
            logger.info(f"Paragraph: {para.text}")
        logger.info("Tables after replacement:")
        for table in doc.tables:
            for row in table.rows:
                row_text = [cell.text.strip() for cell in row.cells]
                logger.info(f"Row: {row_text}")

        return doc
    except Exception as e:
        logger.error(f"Error replacing placeholders: {str(e)}")
        raise